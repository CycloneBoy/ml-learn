#!/user/bin/env python
# -*- coding: utf-8 -*-
# @File  : use_chrome_to_download.py
# @Author: sl
# @Date  : 2022/2/12 - 下午2:52
import os
import time
import traceback
from collections import OrderedDict
from typing import List

from docx import Document
from docx.shared import Inches

from fpdf import FPDF
from selenium import webdriver
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.support.wait import WebDriverWait

import time
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options

from test.pdf.extract_pdf_word import extract_words_from_txt_and_to_dict
from util.file_utils import save_to_json, load_to_json, get_file_name, list_file, read_to_text_list
from util.logger_utils import logger


def atest_chrome2(url):
    b = webdriver.Chrome()
    wait = WebDriverWait(b, 60)
    b.get(url)
    time.sleep(20)
    b.quit()


PATTERN_YOUTUBE_DOWNSUB = "https://downsub.com/?url=https%3A%2F%2Fyoutu.be%2F{}"
DOWNLOAD_DIR = "/home/sl/workspace/data/video/youtube/audio"
GLOBAL_VIDEO_LIST_DICT = f"{DOWNLOAD_DIR}/video_list_dict.json"


class YoutubeSubtitleDownload(object):

    def __init__(self, download_dir=DOWNLOAD_DIR):
        self.download_dir = download_dir
        self.options = Options()
        # service = Service(ChromeDriverManager().install())
        # driver = webdriver.Chrome(service=service, options=options)
        self.driver = None

    def get_driver(self, download_dir=None):
        if download_dir is None:
            download_dir = self.download_dir
        self.options.add_experimental_option('prefs', {
            "download.default_directory": download_dir,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True
        })
        self.driver = webdriver.Chrome(options=self.options)

    def download_one_subtitle(self, src_xpath):
        """
        下载一个字幕
        :param driver:
        :param src_xpath:
        :return:
        """
        try:
            btn_english: WebElement = self.driver.find_element(by=By.XPATH, value=src_xpath)
            btn_english.click()
            # print(f"完成下载：{btn_english.text}")
        except:
            logger.error("出现错误")

    def download_one_language(self, xpath_english):
        """
        下载一个国家的字幕
        :param xpath_english:
        :return:
        """
        xpath_english_srt = xpath_english.format(1)
        xpath_english_txt = xpath_english.format(2)
        self.download_one_subtitle(src_xpath=xpath_english_srt)
        self.download_one_subtitle(src_xpath=xpath_english_txt)

    def get_youtube_subtitle(self, url, download_dir=None):
        """
        下载youtube指定视频的字幕文件

        :param url: https://downsub.com/?url=https%3A%2F%2Fyoutu.be%2FMdUkC7Vz3rg
        :param download_dir:
        :return:
        """
        self.get_driver(download_dir=download_dir)
        self.driver.get(url)

        # time.sleep(3)

        xpath_english = '//*[@id="app"]/div[1]/main/div/div[2]/div/div[1]/div[1]/div[2]/div[1]/button[{}]'
        self.download_one_language(xpath_english)

        xpath_chinese_simplified = '//*[@id="app"]/div/main/div/div[2]/div/div[1]/div[1]/div[2]/div[21]/button[{}]'
        self.download_one_language(xpath_chinese_simplified)

        self.driver.quit()

    def get_url(self):
        self.driver.get(url)

    def close(self):
        self.driver.quit()

    @staticmethod
    def make_url(video_url):
        """
        base YOUTUBE video url to generate downsub url
        :param video_url:  https://youtu.be/l2Z1_wNTmJc
        :return:
        """
        begin_index = str(video_url).rfind("/")
        short_video_id = video_url[begin_index + 1:]
        down_url = PATTERN_YOUTUBE_DOWNSUB.format(short_video_id)
        return down_url

    @staticmethod
    def from_list_url_to_id(video_url):
        """
        根据列表的url 提取视频的ID
        :param video_url:  https://www.youtube.com/watch?v=MdUkC7Vz3rg&list=PLdz6EbLJcjJ9ixS2JC_DDFekyC_jTeVLL&index=1
        :return:
        """
        begin_index = str(video_url).rfind("v=")
        end_index = str(video_url).rfind("&list")
        short_video_id = video_url[begin_index + 2:end_index]
        return short_video_id

    @staticmethod
    def get_video_list_form_catch(url):
        """

        :param url:
        :return:
        """
        if not os.path.exists(GLOBAL_VIDEO_LIST_DICT):
            save_to_json(GLOBAL_VIDEO_LIST_DICT, {"URL": "NAME"})
        video_list_dict = load_to_json(GLOBAL_VIDEO_LIST_DICT)

        video_list_name = None
        if url in video_list_dict:
            video_list_name = video_list_dict[url]
        return video_list_name

    def get_youtube_list(self, url, ):
        """
        下载youtube指定视频列表

        :param url:
        :param download_dir:
        :return:
        """
        video_list_name = YoutubeSubtitleDownload.get_video_list_form_catch(url)
        if video_list_name is not None:
            video_list = load_to_json(video_list_name)
            video_list_dir = get_file_name(video_list_name)[:-5]
            return video_list, video_list_dir

        # to download list first
        self.get_driver()
        # self.driver = webdriver.Chrome(options=self.options)

        self.driver.get(url)

        time.sleep(3)

        video_list_title_xpath = '//*[@id="title"]/yt-formatted-string/a'
        ele_video_list_title: WebElement = self.driver.find_element(by=By.XPATH, value=video_list_title_xpath)
        video_list_title = ele_video_list_title.text

        ele_video_title_list: List[WebElement] = self.driver.find_elements(by=By.ID, value="video-title")
        video_total = len(ele_video_title_list)

        logger.info(f"video list url: {url}")
        logger.info(f"video list total: {video_total}")
        video_list_info = []
        for index, video_title in enumerate(ele_video_title_list):
            title = video_title.text
            video_href = video_title.get_property("href")
            video_short_id = YoutubeSubtitleDownload.from_list_url_to_id(video_href)
            logger.info(f"{index} - {title} - {video_href}")

            video_info = {
                "index": index + 1,
                "title": title,
                "href": video_href,
                "id": video_short_id,
                "subtitle_href": PATTERN_YOUTUBE_DOWNSUB.format(video_short_id)
            }
            video_list_info.append(video_info)

        file_name = f"{self.download_dir}/{video_list_title}_{video_total}.json"
        file_name = file_name.replace(" ", "_")
        save_to_json(file_name, video_list_info)
        video_list_dir = get_file_name(file_name)[:-5]

        video_list_dict = load_to_json(GLOBAL_VIDEO_LIST_DICT)
        video_list_dict[url] = file_name
        save_to_json(GLOBAL_VIDEO_LIST_DICT, video_list_dict)

        self.driver.quit()

        return video_list_info, video_list_dir

    def download_video_list_subtitle(self, url):
        """
        get video list subtitle
        :param url:
        :return:
        """
        video_list_info, video_list_dir = self.get_youtube_list(url)
        download_dir = os.path.join(self.download_dir, video_list_dir)

        for index, video_info in enumerate(video_list_info):
            if self.check_video_subtitle_exist(video_info=video_info, video_list_dir=video_list_dir):
                continue
            try:
                self.get_youtube_subtitle(video_info["subtitle_href"], download_dir=download_dir)
                logger.info(f"success download: {index + 1} - {video_info['title']}")
            except Exception as e:
                traceback.print_exc(e)
                logger.error(f"success failed: {index + 1} - {video_info['title']}")

    def check_video_subtitle_exist(self, video_info, video_list_dir, is_txt=True):
        """
        检查是否已经下载
        :return:
        """
        # [English (auto-generated)] Speak English Like a Native Speaker in 20 Minutes [DownSub.com].srt
        # [Chinese Traditional] Speak English Like a Native Speaker in 20 Minutes [DownSub.com].txt
        target_file_name = self.get_video_subtitle(video_info, is_txt=is_txt)

        download_dir = os.path.join(self.download_dir, video_list_dir)
        all_file_name = list_file(download_dir)
        file_name_set = set(all_file_name)

        flag = False
        if target_file_name in file_name_set:
            flag = True

        return flag

    def get_video_subtitle(self, video_info, is_txt=True):
        """
        获取视频的字幕
        :param video_info:
        :param is_txt:
        :return:
        """
        title = video_info["title"]
        title = title.replace(":", "_")
        end = "txt" if is_txt else "srt"
        target_file_name = f"[English (auto-generated)] {title} [DownSub.com].{end}"

        return target_file_name

    def to_pdf(self, url, is_txt=True):
        video_list_info, video_list_dir = self.get_youtube_list(url)
        download_dir = os.path.join(self.download_dir, video_list_dir)

        video_list_title = video_list_dir
        video_total = len(video_list_info)
        video_subtitle_list = []
        video_title = []
        for index, video_info in enumerate(video_list_info):
            if not self.check_video_subtitle_exist(video_info=video_info, video_list_dir=video_list_dir, is_txt=is_txt):
                continue

            target_file_name = self.get_video_subtitle(video_info, is_txt=is_txt)
            file_name = os.path.join(download_dir, target_file_name)
            video_subtitle = read_to_text_list(file_name)

            video_title.append(target_file_name)
            video_subtitle_list.append(video_subtitle)

        # to docx
        document = Document()
        document.add_heading(video_list_title, 0)

        for index, subtitle in enumerate(video_subtitle_list):
            video_title = video_list_info[index]["title"]
            video_href = video_list_info[index]["href"]

            document.add_heading(video_title, level=1)
            document.add_paragraph(video_href, style='Intense Quote')

            for item in subtitle:
                new_item = str(item).strip("\n")
                if len(new_item) <= 1:
                    continue
                p = document.add_paragraph(new_item)

        file_name = f"{self.download_dir}/{video_list_title}.docx"
        document.save(file_name)

        logger.info(f"output pdf file: {file_name}")

    def to_docx(self, video_title_list):
        """
        字幕保存到 docx
        :param video_title_list:
        :return:
        """
        document = Document()

        document.add_heading('Document Title', 0)

        # pdf = FPDF()
        # pdf.add_page()
        # pdf.set_font("Arial", size=15)
        #
        # # insert the texts in pdf
        # for index, subtitle in enumerate(video_subtitle_list):
        #     video_title = video_list_info[index]["title"]
        #     video_href = video_list_info[index]["href"]
        #
        #     pdf.cell(200, 10, txt=video_title, border=1, ln=1, align='C')
        #     pdf.cell(200, 10, txt=video_href, ln=1, align='C')
        #     for item in subtitle:
        #         pdf.write(item)
        #
        #     pdf.cell(200, 10, txt="[END]", ln=1, align='C')
        #     pdf.cell(200, 10, txt="", ln=1, align='C')
        #     pdf.cell(200, 10, txt="", ln=1, align='C')
        #     # pdf.cell(200, 10, txt=x, ln=1, align='C')

    def upload_dict_to_eudic(self, url, upload_name=None):
        """
        upload video subtitle list to eudic

        :param url:
        :param upload_name:
        :return:
        """
        video_list_info, video_list_dir = self.get_youtube_list(url)
        download_dir = os.path.join(self.download_dir, video_list_dir)

        # to eudic dict
        extract_words_from_txt_and_to_dict(txt_dir=download_dir,
                                           save_dir=f"{download_dir}_dict",
                                           category_name=f"youtube_{video_list_dir}")

    def download_video_list_subtitle_and_to_docx(self, url, ):
        """

        :return:
        """

        self.download_video_list_subtitle(url)
        self.to_pdf(url)


if __name__ == '__main__':
    pass

    video_list = [
        "https://downsub.com/?url=https%3A%2F%2Fyoutu.be%2FMdUkC7Vz3rg",
        "https://downsub.com/?url=https%3A%2F%2Fyoutu.be%2Fl2Z1_wNTmJc",
        ""
    ]
    # url = "https://downsub.com/?url=https%3A%2F%2Fyoutu.be%2FMdUkC7Vz3rg"
    url = "https://www.youtube.com/playlist?list=PLdz6EbLJcjJ9ixS2JC_DDFekyC_jTeVLL"
    # url = 'https://sec.report/Document/0001670254-20-001152/document_1.pdf'

    # atest_chrome2(url)

    title_downloader = YoutubeSubtitleDownload()
    # title_downloader.download_video_list_subtitle_and_to_docx(url)
    title_downloader.upload_dict_to_eudic(url)
    # title_downloader.to_pdf(url)
    # title_downloader.download_video_list_subtitle(url)
    # title_downloader.get_youtube_list(url)
    # title_downloader.chrome_youtube_subtitle(url)

    # # res = YoutubeSubtitleDownload.make_url("https://youtu.be/l2Z1_wNTmJc")
    # res = YoutubeSubtitleDownload.from_list_url_to_id(
    #     "https://www.youtube.com/watch?v=MdUkC7Vz3rg&list=PLdz6EbLJcjJ9ixS2JC_DDFekyC_jTeVLL&index=1")
    # print(res)
